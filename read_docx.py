"""
Script to extract text content from Word (.docx) files
"""
import sys

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx library not installed")
    print("Please install it using: pip install python-docx")
    sys.exit(1)

def read_docx(file_path):
    """Extract all text from a Word document"""
    try:
        doc = Document(file_path)
        
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except Exception as e:
        return f"ERROR: Failed to read document - {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <path_to_docx_file>")
        sys.exit(1)
    
    # Set UTF-8 encoding for output
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    
    file_path = sys.argv[1]
    content = read_docx(file_path)
    print(content)
